import logging
import os
import re
from datetime import datetime, timezone

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Regex for inline bold, italic, hyperlinks, and bare URLs
_INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*"                          # **bold**
    r"|\*(.+?)\*"                              # *italic*
    r"|\[([^\]]+)\]\(([^)]+)\)"                # [text](url)
    r"|(?<![(\w])(https?://[^\s<>\")]+[^\s<>\".,;:!?)_])"  # bare URL
    r")"
)


def _sanitize_filename(title):
    """Remove characters unsafe for filenames."""
    safe = re.sub(r'[<>:"/\\|?*]', "", title)
    safe = safe.strip(". ")
    return safe[:200] if safe else "untitled"


def _add_hyperlink(paragraph, url, text, bold=False, italic=False):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Explicit blue color + underline (don't rely on Hyperlink style existing)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:i")
        rPr.append(i)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_formatted_runs(paragraph, text, bold=False, italic=False):
    """Parse inline **bold**, *italic*, and [text](url) markers and add runs."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        if m.group(2) is not None:
            # **bold** — recurse to handle nested links/italic
            _add_formatted_runs(paragraph, m.group(2), bold=True, italic=italic)
        elif m.group(3) is not None:
            # *italic* — recurse to handle nested links/bold
            _add_formatted_runs(paragraph, m.group(3), bold=bold, italic=True)
        elif m.group(4) is not None:
            # [text](url)
            _add_hyperlink(paragraph, m.group(5), m.group(4), bold=bold, italic=italic)
        elif m.group(6) is not None:
            # bare URL
            _add_hyperlink(paragraph, m.group(6), m.group(6), bold=bold, italic=italic)
        pos = m.end()
    # Remaining plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def _add_horizontal_rule(doc):
    """Add a horizontal rule as an empty paragraph with a bottom border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _flush_table(doc, table_lines):
    """Parse accumulated markdown table lines and add a docx table."""
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows like |---|---|
        if all(re.fullmatch(r":?-+:?", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""  # clear default paragraph
            p = cell.paragraphs[0]
            if i == 0:
                # Bold header row
                run = p.add_run(cell_text)
                run.bold = True
            else:
                _add_formatted_runs(p, cell_text)


_NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def create_document(title, content, output_dir, timestamp=""):
    """Create a .docx file from markdown content.

    Returns dict with filename and path.
    """
    safe_title = _sanitize_filename(title)
    if timestamp:
        filename = f"{safe_title}_{timestamp}.docx"
    else:
        filename = f"{safe_title}.docx"
    filepath = os.path.join(output_dir, filename)

    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    doc.add_heading(safe_title, level=0)

    table_buf = []

    for line in content.split("\n"):
        stripped = line.strip()

        # Accumulate table rows
        if stripped.startswith("|"):
            table_buf.append(stripped)
            continue
        # Flush any accumulated table when we hit a non-table line
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []

        if not stripped:
            continue
        if stripped.startswith("### "):
            p = doc.add_heading("", level=3)
            _add_formatted_runs(p, stripped[4:])
        elif stripped.startswith("## "):
            p = doc.add_heading("", level=2)
            _add_formatted_runs(p, stripped[3:])
        elif stripped.startswith("# "):
            p = doc.add_heading("", level=1)
            _add_formatted_runs(p, stripped[2:])
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, stripped[2:])
        elif _NUMBERED_RE.match(stripped):
            m = _NUMBERED_RE.match(stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, m.group(2))
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            _add_formatted_runs(p, stripped[2:])
        elif re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            _add_horizontal_rule(doc)
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)

    # Flush trailing table
    if table_buf:
        _flush_table(doc, table_buf)

    doc.save(filepath)
    logger.info("Created document: %s", filepath)
    return {"filename": filename, "path": filepath}


def list_documents(output_dir):
    """List existing .docx files in the output directory.

    Returns dict with files list.
    """
    if not os.path.isdir(output_dir):
        return {"files": []}

    files = []
    for name in sorted(os.listdir(output_dir)):
        if not name.endswith(".docx"):
            continue
        path = os.path.join(output_dir, name)
        stat = os.stat(path)
        files.append({
            "name": name,
            "path": path,
            "modified": datetime.fromtimestamp(
                stat.st_mtime, tz=timezone.utc
            ).isoformat(),
        })

    return {"files": files}
